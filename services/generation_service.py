import os
import requests
from typing import Dict, List, Optional
from services.retrieval_service import RetrievalService
from services.persistence_service import PersistenceService

class GenerationService:
    def __init__(self, retrieval_service: RetrievalService, 
                 persistence_service: PersistenceService):
        self.retrieval = retrieval_service
        self.persistence = persistence_service
        self.base_url = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434")
        self.chat_model = os.environ.get("OLLAMA_CHAT_MODEL", "qwen2.5:3b")
        self.generated_dir = "data/generated"
        os.makedirs(self.generated_dir, exist_ok=True)

    def generate(self, prompt: str, mode: str = "prompt_libre",
                 source_doc_ids: List[int] = None,
                 doc_format: str = "markdown",
                 user_id: int = None) -> Dict:
        """
        Genera un documento nuevo usando IA.
        Modos: prompt_libre, basado_repositorio, basado_documento
        """
        context_text = ""
        used_doc_ids = []

        # Recuperar contexto si aplica
        if mode in ("basado_repositorio", "basado_documento"):
            doc_id_filter = source_doc_ids[0] if (
                mode == "basado_documento" and source_doc_ids
            ) else None

            results = self.retrieval.search(
                query=prompt,
                top_k=8,
                document_id=doc_id_filter
            )

            if results:
                context_parts = [r["text"] for r in results]
                context_text = "\n\n".join(context_parts)
                used_doc_ids = list({r["document_id"] for r in results})

        # Construir prompt según modo
        if mode == "prompt_libre":
            full_prompt = f"""Eres un asistente experto en redacción institucional.
Genera un documento profesional basado en las siguientes instrucciones.
El documento debe estar bien estructurado, con secciones claras.
Usa formato Markdown si el usuario no especifica otro formato.

INSTRUCCIONES DEL USUARIO:
{prompt}

DOCUMENTO GENERADO:"""

        else:
            full_prompt = f"""Eres un asistente experto en redacción institucional.
Genera un documento profesional basado en las instrucciones del usuario
y el contenido de referencia proporcionado.
El documento debe estar bien estructurado y fundamentado en el contexto.
Usa formato Markdown si el usuario no especifica otro formato.

CONTENIDO DE REFERENCIA:
{context_text}

INSTRUCCIONES DEL USUARIO:
{prompt}

DOCUMENTO GENERADO:"""

        # Llamar al LLM
        try:
            response = requests.post(
                f"{self.base_url}/api/generate",
                json={"model": self.chat_model, 
                      "prompt": full_prompt, 
                      "stream": False},
                timeout=120
            )
            response.raise_for_status()
            content = response.json().get("response", "").strip()
        except Exception as e:
            return {"success": False, "error": str(e)}

        if not content:
            return {"success": False, "error": "El modelo no generó contenido."}

        # Extraer título de la primera línea
        lines = content.strip().split("\n")
        title = lines[0].replace("#", "").strip()
        if not title or len(title) > 200:
            title = prompt[:100]

        word_count = len(content.split())

        # Persistir en DB
        try:
            conn = self.persistence.db.get_connection()
            cur = conn.cursor()
            cur.execute("""
                INSERT INTO generated_documents_v2
                (user_id, title, prompt, content, format, 
                 generation_mode, source_doc_ids, model_used, word_count)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id;
            """, (
                user_id, title, prompt, content, doc_format,
                mode, source_doc_ids or [], 
                self.chat_model, word_count
            ))
            gen_id = cur.fetchone()[0]
            conn.commit()
            conn.close()
        except Exception as e:
            print(f"[ERROR] No se pudo persistir el documento generado: {e}")
            gen_id = None

        return {
            "success": True,
            "id": gen_id,
            "title": title,
            "content": content,
            "word_count": word_count,
            "mode": mode,
            "used_doc_ids": used_doc_ids
        }

    def get_all(self, user_id: int = None) -> List[Dict]:
        """Lista todos los documentos generados."""
        conn = self.persistence.db.get_connection()
        from psycopg2.extras import RealDictCursor
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                if user_id:
                    cur.execute("""
                        SELECT id, title, format, generation_mode, 
                               word_count, created_at
                        FROM generated_documents_v2
                        WHERE user_id = %s
                        ORDER BY created_at DESC;
                    """, (user_id,))
                else:
                    cur.execute("""
                        SELECT id, title, format, generation_mode,
                               word_count, created_at
                        FROM generated_documents_v2
                        ORDER BY created_at DESC;
                    """)
                return [dict(r) for r in cur.fetchall()]
        finally:
            conn.close()

    def get_by_id(self, gen_id: int) -> Optional[Dict]:
        """Obtiene un documento generado por ID."""
        conn = self.persistence.db.get_connection()
        from psycopg2.extras import RealDictCursor
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute(
                    "SELECT * FROM generated_documents_v2 WHERE id = %s;",
                    (gen_id,)
                )
                row = cur.fetchone()
                return dict(row) if row else None
        finally:
            conn.close()

    def delete(self, gen_id: int) -> bool:
        """Elimina un documento generado."""
        conn = self.persistence.db.get_connection()
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "DELETE FROM generated_documents_v2 WHERE id = %s;",
                    (gen_id,)
                )
            conn.commit()
            return True
        except Exception:
            return False
        finally:
            conn.close()

    def export_docx(self, gen_id: int) -> bytes:
        """Exporta documento generado como DOCX en memoria."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from io import BytesIO
        import re

        doc_data = self.get_by_id(gen_id)
        if not doc_data:
            raise ValueError("Documento no encontrado")

        doc = Document()

        # Estilo del documento
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Procesar contenido línea por línea
        for line in doc_data["content"].split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            if line.startswith("# "):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                doc.add_paragraph(re.sub(r'^\d+\. ', '', line), 
                                  style='List Number')
            else:
                # Limpiar markdown bold/italic
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                clean = re.sub(r'\*(.*?)\*', r'\1', clean)
                doc.add_paragraph(clean)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def export_pdf(self, gen_id: int) -> bytes:
        """Exporta documento generado como PDF en memoria."""
        from fpdf import FPDF
        from io import BytesIO
        import re

        doc_data = self.get_by_id(gen_id)
        if not doc_data:
            raise ValueError("Documento no encontrado")

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_margins(20, 20, 20)

        for line in doc_data["content"].split("\n"):
            line = line.strip()
            if not line:
                pdf.ln(4)
                continue

            if line.startswith("# "):
                pdf.set_font("Helvetica", "B", 18)
                pdf.set_text_color(30, 30, 30)
                clean = line[2:]
                pdf.multi_cell(0, 10, clean)
                pdf.ln(2)
            elif line.startswith("## "):
                pdf.set_font("Helvetica", "B", 14)
                pdf.set_text_color(50, 50, 50)
                clean = line[3:]
                pdf.multi_cell(0, 8, clean)
                pdf.ln(1)
            elif line.startswith("### "):
                pdf.set_font("Helvetica", "B", 12)
                pdf.set_text_color(70, 70, 70)
                clean = line[4:]
                pdf.multi_cell(0, 7, clean)
            elif line.startswith("- ") or line.startswith("* "):
                pdf.set_font("Helvetica", "", 11)
                pdf.set_text_color(40, 40, 40)
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line[2:])
                clean = re.sub(r'\*(.*?)\*', r'\1', clean)
                pdf.multi_cell(0, 6, f"  • {clean}")
            else:
                pdf.set_font("Helvetica", "", 11)
                pdf.set_text_color(40, 40, 40)
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                clean = re.sub(r'\*(.*?)\*', r'\1', clean)
                try:
                    pdf.multi_cell(0, 6, clean)
                except:
                    pdf.multi_cell(0, 6, clean.encode('latin-1', 
                                   errors='replace').decode('latin-1'))

        return bytes(pdf.output())
